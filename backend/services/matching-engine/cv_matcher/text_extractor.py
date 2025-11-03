"""
TextExtractor - Extract text from CV files (PDF, DOCX)
=======================================================

Handles text extraction from various CV formats.
"""

import logging
from typing import Optional
import os

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text content from CV files"""

    def __init__(self):
        """Initialize text extractor with required libraries"""
        self._pdf_extractor = None
        self._docx_extractor = None

    def extract(self, file_path: str) -> str:
        """
        Extract text from CV file.

        Args:
            file_path: Path to CV file (PDF, DOCX, DOC, or TXT)

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format not supported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"CV file not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        logger.info("Extracting text from %s", file_path)

        if file_ext == ".pdf":
            text = self._extract_from_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            text = self._extract_from_docx(file_path)
        elif file_ext == ".txt":
            text = self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        logger.info("Extracted %d characters from %s", len(text), file_path)
        return text

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2

            text = []
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text.append(page.extract_text())

            return "\n".join(text)

        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise
        except Exception as e:
            logger.error("Failed to extract PDF: %s", str(e))
            raise

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            import docx

            doc = docx.Document(file_path)
            text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)

            return "\n".join(text)

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise
        except Exception as e:
            logger.error("Failed to extract DOCX: %s", str(e))
            raise

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()
        except Exception as e:
            logger.error("Failed to extract TXT: %s", str(e))
            raise
